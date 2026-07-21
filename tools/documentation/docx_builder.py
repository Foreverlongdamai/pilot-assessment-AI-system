"""Deterministic Markdown-to-DOCX renderer for Pilot Assessment manuals."""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from manual_common import (
    MANUAL_ROOT,
    REPOSITORY_ROOT,
    DocumentationError,
    read_json,
    replace_document_references,
)
from markdown_it import MarkdownIt
from markdown_it.token import Token
from PIL import Image

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120
NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5A6470"
GREEN = "1F7A64"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F4EF"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
CODE_FONT = "Cascadia Mono"
ASSET_ONLY = re.compile(r"^\[\[ASSET:([a-z0-9][a-z0-9-]*)\]\]$")
DOCUMENT_TIMESTAMP = datetime(2026, 7, 21, 0, 0, 0)


@dataclass(frozen=True)
class HeadingEntry:
    level: int
    text: str
    bookmark: str


@dataclass
class ListState:
    kind: str
    counter: int = 0
    marker_pending: bool = False


def _set_run_font(
    run: Any,
    *,
    name: str = BODY_FONT,
    east_asia: str = CJK_FONT,
    size: float | None = None,
    colour: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if colour is not None:
        run.font.color.rgb = RGBColor.from_string(colour)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style: Any, *, name: str, east_asia: str, size: float, colour: str) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(colour)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _shade(element: Any, fill: str) -> None:
    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _paragraph_shading(paragraph: Any, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _paragraph_left_border(paragraph: Any, colour: str, size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), colour)
    borders.append(left)


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def _normalised_column_widths(rows: list[list[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    weights: list[float] = []
    for column in range(columns):
        lengths = [len(row[column]) if column < len(row) else 0 for row in rows]
        weights.append(float(max(10, min(max(lengths, default=10), 55))))
    total = sum(weights)
    widths = [max(900, round(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 900:
        shortage = 900 - widths[-1]
        widths[-1] = 900
        widest = max(range(len(widths) - 1), key=widths.__getitem__)
        widths[widest] -= shortage
    return widths


def _apply_table_geometry(table: Any, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_hyperlink(
    paragraph: Any, text: str, *, target: str | None = None, anchor: str | None = None
) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    elif target:
        relationship_id = paragraph.part.relate_to(
            target,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink.set(qn("r:id"), relationship_id)
    else:
        raise DocumentationError("hyperlink requires a target or anchor")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([colour, underline])
    run.append(run_properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_field(paragraph: Any, instruction: str, placeholder: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend([begin, code, separate, text, end])
    paragraph._p.append(run)


def _set_update_fields(document: DocumentType) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def _configure_styles(document: DocumentType) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, name=BODY_FONT, east_asia=CJK_FONT, size=11, colour=NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, colour, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
        ("Heading 4", 11, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        _set_style_font(style, name=BODY_FONT, east_asia=CJK_FONT, size=size, colour=colour)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        _set_style_font(style, name=BODY_FONT, east_asia=CJK_FONT, size=11, colour=NAVY)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "PA Code Block" not in styles:
        code_style = styles.add_style("PA Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["PA Code Block"]
    _set_style_font(code_style, name=CODE_FONT, east_asia=CJK_FONT, size=9, colour=NAVY)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.05

    if "PA Caption" not in styles:
        caption_style = styles.add_style("PA Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption_style = styles["PA Caption"]
    _set_style_font(caption_style, name=BODY_FONT, east_asia=CJK_FONT, size=9, colour=MUTED)
    caption_style.font.italic = True
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(4)
    caption_style.paragraph_format.space_after = Pt(9)
    caption_style.paragraph_format.keep_with_next = True

    if "PA TOC" not in styles:
        toc_style = styles.add_style("PA TOC", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc_style = styles["PA TOC"]
    _set_style_font(toc_style, name=BODY_FONT, east_asia=CJK_FONT, size=10.5, colour=BLUE)
    toc_style.paragraph_format.space_after = Pt(3)


def _configure_section(document: DocumentType) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def create_reference_template(path: Path) -> None:
    document = Document()
    _configure_section(document)
    _configure_styles(document)
    _set_update_fields(document)
    document.core_properties.title = "Pilot Assessment System Reference Manual Template"
    document.core_properties.subject = "M8C style-only DOCX reference template"
    document.core_properties.keywords = "Pilot Assessment System, M8C, documentation template"
    document.core_properties.created = DOCUMENT_TIMESTAMP
    document.core_properties.modified = DOCUMENT_TIMESTAMP
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


class ManualDocxBuilder:
    def __init__(
        self,
        *,
        template: Path,
        catalog: dict[str, Any],
        document_record: dict[str, Any],
        language: str,
        metadata: dict[str, Any],
        source_path: Path,
    ) -> None:
        self.document = Document(template)
        self.catalog = catalog
        self.document_record = document_record
        self.language = language
        self.metadata = metadata
        self.source_path = source_path
        self._bookmark_id = 10
        self._heading_cursor = 0
        self.diagram_manifest = read_json(MANUAL_ROOT / "assets" / "diagrams" / "manifest.json")
        self.diagram_index = {
            str(item["asset_id"]): item for item in self.diagram_manifest["diagrams"]
        }
        _configure_section(self.document)
        _configure_styles(self.document)
        _set_update_fields(self.document)
        for paragraph in list(self.document.paragraphs):
            if not paragraph.text:
                paragraph._element.getparent().remove(paragraph._element)
        self._configure_properties()
        self._configure_header_footer()

    def _configure_properties(self) -> None:
        properties = self.document.core_properties
        properties.title = str(self.metadata["title"])
        properties.subject = str(self.metadata["scope"])
        properties.author = "Pilot Assessment System"
        properties.keywords = ", ".join(
            [
                str(self.metadata["document_id"]),
                str(self.metadata["language"]),
                str(self.metadata["scientific_status"]),
            ]
        )
        properties.comments = "Generated from authoritative Markdown by the M8C toolchain."
        properties.created = DOCUMENT_TIMESTAMP
        properties.modified = DOCUMENT_TIMESTAMP

    def _configure_header_footer(self) -> None:
        section = self.document.sections[0]
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run(f"PILOT ASSESSMENT SYSTEM   |   {self.metadata['short_title']}")
        _set_run_font(run, size=8.5, colour=MUTED, bold=True)
        p_pr = header._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), "D7DBE2")
        borders.append(bottom)
        p_pr.append(borders)

        footer = section.footer.paragraphs[0]
        footer.paragraph_format.space_before = Pt(0)
        footer.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT
        )
        left = footer.add_run(
            f"{self.metadata['document_id']}  |  v{self.metadata['document_version']}"
        )
        _set_run_font(left, size=8, colour=MUTED)
        footer.add_run("\t")
        page_label = "第 " if self.language == "zh-CN" else "Page "
        of_label = " 页，共 " if self.language == "zh-CN" else " of "
        page_run = footer.add_run(page_label)
        _set_run_font(page_run, size=8, colour=MUTED)
        _add_field(footer, "PAGE", "1")
        middle = footer.add_run(of_label)
        _set_run_font(middle, size=8, colour=MUTED)
        _add_field(footer, "NUMPAGES", "1")

    def _add_cover(self) -> None:
        icon = (
            REPOSITORY_ROOT
            / "src"
            / "PilotAssessment.Desktop"
            / "Assets"
            / "Brand"
            / "PilotAssessmentIcon-1024.png"
        )
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(34)
        paragraph.paragraph_format.space_after = Pt(16)
        if icon.is_file():
            paragraph.add_run().add_picture(str(icon), width=Inches(0.8))

        kicker = self.document.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kicker.paragraph_format.space_after = Pt(14)
        run = kicker.add_run("PILOT ASSESSMENT SYSTEM")
        _set_run_font(run, size=10, colour=GREEN, bold=True)

        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(8)
        title.paragraph_format.keep_with_next = True
        run = title.add_run(str(self.metadata["title"]))
        _set_run_font(run, size=28, colour=NAVY, bold=True)

        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(26)
        run = subtitle.add_run(
            "正式工程手册" if self.language == "zh-CN" else "Engineering Product Manual"
        )
        _set_run_font(run, size=13, colour=BLUE)

        labels = (
            [
                ("文档 ID", self.metadata["document_id"]),
                ("产品版本", self.metadata["product_version"]),
                ("文档版本", self.metadata["document_version"]),
                ("状态", self.metadata["status"]),
                ("目标读者", ", ".join(self.metadata["audience"])),
                ("科学状态", self.metadata["scientific_status"]),
            ]
            if self.language == "zh-CN"
            else [
                ("Document ID", self.metadata["document_id"]),
                ("Product version", self.metadata["product_version"]),
                ("Document version", self.metadata["document_version"]),
                ("Status", self.metadata["status"]),
                ("Audience", ", ".join(self.metadata["audience"])),
                ("Scientific status", self.metadata["scientific_status"]),
            ]
        )
        for label, value in labels:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(3)
            label_run = paragraph.add_run(f"{label}: ")
            _set_run_font(label_run, size=9.5, colour=MUTED, bold=True)
            value_run = paragraph.add_run(str(value))
            _set_run_font(value_run, size=9.5, colour=NAVY)

        boundary = self.document.add_paragraph()
        boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        boundary.paragraph_format.space_before = Pt(22)
        boundary.paragraph_format.space_after = Pt(0)
        boundary.paragraph_format.left_indent = Inches(0.65)
        boundary.paragraph_format.right_indent = Inches(0.65)
        _paragraph_shading(boundary, LIGHT_GREEN)
        text = (
            "本手册说明工程系统；starter Evidence、BN 与 CPT 尚未获得领域科学校准。"
            if self.language == "zh-CN"
            else (
                "This manual describes the engineering system; starter Evidence, "
                "BN and CPT content is not scientifically calibrated."
            )
        )
        run = boundary.add_run(text)
        _set_run_font(run, size=9.5, colour=NAVY, bold=True)
        self.document.add_page_break()

    def _headings(self, tokens: list[Token]) -> list[HeadingEntry]:
        headings: list[HeadingEntry] = []
        index = 0
        while index < len(tokens):
            token = tokens[index]
            if token.type == "heading_open" and index + 1 < len(tokens):
                inline = tokens[index + 1]
                source_level = int(token.tag[1:])
                if source_level == 1 and inline.content.strip() == str(self.metadata["title"]):
                    index += 3
                    continue
                level = max(1, min(source_level - 1, 4))
                headings.append(
                    HeadingEntry(
                        level=level,
                        text=inline.content.strip(),
                        bookmark=f"pa_heading_{len(headings) + 1:03d}",
                    )
                )
            index += 1
        return headings

    def _add_toc(self, headings: list[HeadingEntry]) -> None:
        title = self.document.add_paragraph()
        title.paragraph_format.space_after = Pt(14)
        run = title.add_run("目录" if self.language == "zh-CN" else "Contents")
        _set_run_font(run, size=20, colour=NAVY, bold=True)
        _add_bookmark(title, "pa_toc", 1)
        for heading in headings:
            paragraph = self.document.add_paragraph(style="PA TOC")
            paragraph.paragraph_format.left_indent = Inches(0.22 * (heading.level - 1))
            _add_hyperlink(paragraph, heading.text, anchor=heading.bookmark)
        note = self.document.add_paragraph()
        note.paragraph_format.space_before = Pt(10)
        text = (
            "目录为确定性内部链接；在 Word 中也可以使用导航窗格浏览真实 Heading 层级。"
            if self.language == "zh-CN"
            else (
                "This deterministic table of contents uses internal links; Word's "
                "Navigation pane also follows the real Heading hierarchy."
            )
        )
        run = note.add_run(text)
        _set_run_font(run, size=8.5, colour=MUTED, italic=True)
        self.document.add_page_break()

    def _render_inline(self, paragraph: Any, token: Token) -> None:
        bold_depth = 0
        italic_depth = 0
        link_target: str | None = None
        link_text: list[str] = []
        children = token.children or []
        for child in children:
            if child.type == "strong_open":
                bold_depth += 1
            elif child.type == "strong_close":
                bold_depth = max(0, bold_depth - 1)
            elif child.type == "em_open":
                italic_depth += 1
            elif child.type == "em_close":
                italic_depth = max(0, italic_depth - 1)
            elif child.type == "link_open":
                link_target = child.attrGet("href")
                link_text = []
            elif child.type == "link_close":
                if link_target:
                    _add_hyperlink(paragraph, "".join(link_text), target=link_target)
                link_target = None
                link_text = []
            elif child.type in {"softbreak", "hardbreak"}:
                if link_target:
                    link_text.append(" ")
                else:
                    paragraph.add_run().add_break()
            elif child.type in {"text", "code_inline"}:
                if link_target:
                    link_text.append(child.content)
                    continue
                run = paragraph.add_run(child.content)
                _set_run_font(
                    run,
                    name=CODE_FONT if child.type == "code_inline" else BODY_FONT,
                    size=9.5 if child.type == "code_inline" else 11,
                    colour=NAVY,
                    bold=bold_depth > 0,
                    italic=italic_depth > 0,
                )
                if child.type == "code_inline":
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), LIGHT_GRAY)
                    run._element.get_or_add_rPr().append(shading)
            elif child.type == "image":
                source = child.attrGet("src")
                if not source:
                    raise DocumentationError("Markdown image is missing src")
                image_path = (self.source_path.parent / source).resolve()
                if not image_path.is_file():
                    raise DocumentationError(f"Markdown image is missing: {image_path}")
                paragraph.add_run().add_picture(str(image_path), width=Inches(6.1))
            else:
                raise DocumentationError(f"unsupported inline Markdown token: {child.type}")

    def _add_asset(self, asset_id: str) -> None:
        item = self.diagram_index.get(asset_id)
        if not item or item.get("status") != "rendered":
            raise DocumentationError(f"diagram is not rendered: {asset_id}")
        png = MANUAL_ROOT / "assets" / "diagrams" / str(item["png"])
        if not png.is_file():
            raise DocumentationError(f"diagram PNG is missing: {png}")
        with Image.open(png) as image:
            width_px, height_px = image.size
        max_width = 6.15
        max_height = 4.65
        ratio = width_px / height_px
        width = min(max_width, max_height * ratio)
        height = width / ratio
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        shape = paragraph.add_run().add_picture(
            str(png), width=Inches(width), height=Inches(height)
        )
        alt_text = str(item["alt_text"][self.language])
        shape._inline.docPr.set("descr", alt_text)
        caption = self.document.add_paragraph(style="PA Caption")
        run = caption.add_run(str(item["caption"][self.language]))
        _set_run_font(run, size=9, colour=MUTED, italic=True)
        explanation = self.document.add_paragraph()
        explanation.paragraph_format.space_after = Pt(8)
        run = explanation.add_run(alt_text)
        _set_run_font(run, size=9, colour=MUTED)

    def _add_table(self, rows: list[list[str]]) -> None:
        if not rows:
            return
        columns = max(len(row) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=columns)
        widths = _normalised_column_widths(rows)
        _apply_table_geometry(table, widths)
        _set_repeat_table_header(table.rows[0])
        for row_index, values in enumerate(rows):
            for column_index in range(columns):
                cell = table.cell(row_index, column_index)
                value = values[column_index] if column_index < len(values) else ""
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                inline_tokens = MarkdownIt(
                    "commonmark",
                    {"html": False, "linkify": False, "typographer": False},
                ).parseInline(value)
                if inline_tokens:
                    self._render_inline(paragraph, inline_tokens[0])
                for run in paragraph.runs:
                    run.font.size = Pt(9.25)
                    if row_index == 0:
                        run.font.bold = True
                if row_index == 0:
                    _shade(cell._tc, LIGHT_BLUE)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _consume_table(tokens: list[Token], start: int) -> tuple[list[list[str]], int]:
        rows: list[list[str]] = []
        current_row: list[str] | None = None
        index = start + 1
        while index < len(tokens):
            token = tokens[index]
            if token.type == "table_close":
                return rows, index + 1
            if token.type == "tr_open":
                current_row = []
            elif token.type == "tr_close":
                if current_row is not None:
                    rows.append(current_row)
                current_row = None
            elif token.type == "inline" and current_row is not None:
                current_row.append(token.content.strip())
            index += 1
        raise DocumentationError("Markdown table is not closed")

    def _add_horizontal_rule(self) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(8)
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D7DBE2")
        borders.append(bottom)
        p_pr.append(borders)

    def _render_body(self, tokens: list[Token], headings: list[HeadingEntry]) -> None:
        index = 0
        list_stack: list[ListState] = []
        blockquote_depth = 0
        while index < len(tokens):
            token = tokens[index]
            token_type = token.type
            if token_type == "heading_open":
                inline = tokens[index + 1]
                source_level = int(token.tag[1:])
                if source_level == 1 and inline.content.strip() == str(self.metadata["title"]):
                    index += 3
                    continue
                heading = headings[self._heading_cursor]
                self._heading_cursor += 1
                paragraph = self.document.add_paragraph(style=f"Heading {heading.level}")
                self._render_inline(paragraph, inline)
                _add_bookmark(paragraph, heading.bookmark, self._bookmark_id)
                self._bookmark_id += 1
                back = paragraph.add_run("  ")
                _set_run_font(back, size=7, colour=MUTED)
                _add_hyperlink(
                    paragraph,
                    "返回目录" if self.language == "zh-CN" else "Back to contents",
                    anchor="pa_toc",
                )
                index += 3
                continue
            if token_type == "table_open":
                rows, index = self._consume_table(tokens, index)
                self._add_table(rows)
                continue
            if token_type in {"bullet_list_open", "ordered_list_open"}:
                if token_type == "ordered_list_open":
                    list_stack.append(
                        ListState(
                            kind="number",
                            counter=int(token.attrGet("start") or 1) - 1,
                        )
                    )
                else:
                    list_stack.append(ListState(kind="bullet"))
            elif token_type in {"bullet_list_close", "ordered_list_close"}:
                list_stack.pop()
            elif token_type == "list_item_open" and list_stack:
                state = list_stack[-1]
                if state.kind == "number":
                    state.counter += 1
                    state.marker_pending = True
            elif token_type == "blockquote_open":
                blockquote_depth += 1
            elif token_type == "blockquote_close":
                blockquote_depth = max(0, blockquote_depth - 1)
            elif token_type == "paragraph_open":
                inline = tokens[index + 1]
                asset_match = ASSET_ONLY.fullmatch(inline.content.strip())
                if asset_match:
                    self._add_asset(asset_match.group(1))
                    index += 3
                    continue
                if list_stack:
                    state = list_stack[-1]
                    paragraph = self.document.add_paragraph(
                        style="List Bullet" if state.kind == "bullet" else None
                    )
                    extra_indent = 0.25 * max(0, len(list_stack) - 1)
                    paragraph.paragraph_format.left_indent = Inches(0.375 + extra_indent)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
                    if state.kind == "number" and state.marker_pending:
                        marker = paragraph.add_run(f"{state.counter}. ")
                        _set_run_font(marker, size=11, colour=NAVY)
                        state.marker_pending = False
                else:
                    paragraph = self.document.add_paragraph()
                if blockquote_depth:
                    paragraph.paragraph_format.left_indent = Inches(0.22)
                    paragraph.paragraph_format.right_indent = Inches(0.12)
                    _paragraph_shading(paragraph, LIGHT_GRAY)
                    _paragraph_left_border(paragraph, BLUE)
                self._render_inline(paragraph, inline)
                index += 3
                continue
            elif token_type in {"fence", "code_block"}:
                paragraph = self.document.add_paragraph(style="PA Code Block")
                _paragraph_shading(paragraph, LIGHT_GRAY)
                _paragraph_left_border(paragraph, GREEN, size="12")
                lines = token.content.rstrip("\n").splitlines() or [""]
                for line_index, line in enumerate(lines):
                    run = paragraph.add_run(line)
                    _set_run_font(run, name=CODE_FONT, size=9, colour=NAVY)
                    if line_index < len(lines) - 1:
                        run.add_break()
            elif token_type == "hr":
                self._add_horizontal_rule()
            elif token_type in {
                "paragraph_close",
                "heading_close",
                "inline",
                "list_item_open",
                "list_item_close",
                "thead_open",
                "thead_close",
                "tbody_open",
                "tbody_close",
                "tr_open",
                "tr_close",
                "th_open",
                "th_close",
                "td_open",
                "td_close",
            }:
                pass
            else:
                raise DocumentationError(f"unsupported block Markdown token: {token_type}")
            index += 1

    def build(self, body: str, output: Path) -> dict[str, Any]:
        body = replace_document_references(body, catalog=self.catalog, language=self.language)
        parser = MarkdownIt("commonmark", {"html": False, "linkify": False, "typographer": False})
        parser.enable("table")
        tokens = parser.parse(body)
        headings = self._headings(tokens)
        if not headings:
            raise DocumentationError(
                f"manual has no buildable section headings: {self.source_path}"
            )
        self._add_cover()
        self._add_toc(headings)
        self._render_body(tokens, headings)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(output)
        return {
            "heading_count": len(headings),
            "paragraph_count": len(self.document.paragraphs),
            "table_count": len(self.document.tables),
            "inline_shape_count": len(self.document.inline_shapes),
        }


def build_manual_docx(
    *,
    template: Path,
    catalog: dict[str, Any],
    document_record: dict[str, Any],
    language: str,
    metadata: dict[str, Any],
    source_path: Path,
    body: str,
    output: Path,
) -> dict[str, Any]:
    builder = ManualDocxBuilder(
        template=template,
        catalog=catalog,
        document_record=document_record,
        language=language,
        metadata=metadata,
        source_path=source_path,
    )
    return builder.build(body, output)


__all__ = ["build_manual_docx", "create_reference_template"]
