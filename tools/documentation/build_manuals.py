"""Build versioned Pilot Assessment DOCX manuals from authoritative Markdown."""

from __future__ import annotations

import argparse
import hashlib
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx_builder import build_manual_docx, create_reference_template
from manual_common import (
    CATALOG_PATH,
    MANUAL_ROOT,
    TOOL_ROOT,
    DocumentationError,
    add_selection_arguments,
    aggregate_manual_source,
    load_catalog,
    output_root,
    parse_manual_source,
    result_payload,
    safe_manual_path,
    selected_variants,
    sha256_file,
    write_json,
)
from validate_manuals import validate

FIXED_ZIP_TIME = (2026, 7, 21, 0, 0, 0)
TEMPLATE_PATH = MANUAL_ROOT / "template" / "pilot-assessment-reference.docx"


def _normalise_docx(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as archive:
        members = [(item.filename, archive.read(item.filename)) for item in archive.infolist()]
    with tempfile.NamedTemporaryFile(
        prefix=f"{path.stem}-", suffix=".docx", dir=path.parent, delete=False
    ) as temporary:
        temporary_path = Path(temporary.name)
    try:
        with zipfile.ZipFile(
            temporary_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
        ) as archive:
            for filename, payload in sorted(members):
                item = zipfile.ZipInfo(filename=filename, date_time=FIXED_ZIP_TIME)
                item.compress_type = zipfile.ZIP_DEFLATED
                item.create_system = 0
                item.external_attr = 0o600 << 16
                archive.writestr(item, payload)
        temporary_path.replace(path)
    finally:
        temporary_path.unlink(missing_ok=True)


def _attribute_int(element: Any, name: str) -> int:
    raw = element.get(qn(name))
    if raw is None:
        raise DocumentationError(f"DOCX geometry attribute is missing: {name}")
    return int(raw)


def _audit_docx(path: Path, *, expected_title: str) -> dict[str, Any]:
    if not zipfile.is_zipfile(path):
        raise DocumentationError(f"generated DOCX is not a valid OOXML ZIP: {path}")
    with zipfile.ZipFile(path, "r") as archive:
        broken = archive.testzip()
        if broken:
            raise DocumentationError(f"generated DOCX has a corrupt member {broken}: {path}")
        document_xml = archive.read("word/document.xml")
        if b"[[DOC:" in document_xml or b"[[ASSET:" in document_xml:
            raise DocumentationError(
                f"generated DOCX contains unresolved stable references: {path}"
            )
        if b"w:bookmarkStart" not in document_xml or b"w:hyperlink" not in document_xml:
            raise DocumentationError(f"generated DOCX is missing static navigation: {path}")
    document = Document(path)
    if document.core_properties.title != expected_title:
        raise DocumentationError(f"generated DOCX title metadata mismatch: {path}")
    if not document.sections:
        raise DocumentationError(f"generated DOCX has no section: {path}")
    section = document.sections[0]
    expected_page = (8.5, 11.0)
    actual_page = (section.page_width.inches, section.page_height.inches)
    if any(
        abs(actual - expected) > 0.01
        for actual, expected in zip(actual_page, expected_page, strict=True)
    ):
        raise DocumentationError(f"generated DOCX is not US Letter portrait: {path}")
    for margin in (
        section.top_margin,
        section.right_margin,
        section.bottom_margin,
        section.left_margin,
    ):
        if abs(margin.inches - 1.0) > 0.01:
            raise DocumentationError(f"generated DOCX margin is not 1 inch: {path}")
    heading_count = len(
        [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name.startswith("Heading ")
        ]
    )
    if heading_count == 0:
        raise DocumentationError(f"generated DOCX has no real Heading styles: {path}")
    for table in document.tables:
        table_pr = table._tbl.tblPr
        table_width = table_pr.find(qn("w:tblW"))
        table_indent = table_pr.find(qn("w:tblInd"))
        table_layout = table_pr.find(qn("w:tblLayout"))
        if table_width is None or _attribute_int(table_width, "w:w") != 9360:
            raise DocumentationError(f"generated DOCX table width is not 9360 DXA: {path}")
        if table_indent is None or _attribute_int(table_indent, "w:w") != 120:
            raise DocumentationError(f"generated DOCX table indent is not 120 DXA: {path}")
        if table_layout is None or table_layout.get(qn("w:type")) != "fixed":
            raise DocumentationError(f"generated DOCX table layout is not fixed: {path}")
        grid_widths = [_attribute_int(column, "w:w") for column in list(table._tbl.tblGrid)]
        if sum(grid_widths) != 9360:
            raise DocumentationError(f"generated DOCX table grid does not sum to 9360: {path}")
        for row in table.rows:
            cell_widths = [
                _attribute_int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")), "w:w")
                for cell in row.cells
            ]
            if cell_widths != grid_widths:
                raise DocumentationError(f"generated DOCX cell/grid width mismatch: {path}")
    return {
        "heading_count": heading_count,
        "table_count": len(document.tables),
        "paragraph_count": len(document.paragraphs),
        "inline_shape_count": len(document.inline_shapes),
    }


def build(*, status: str, language: str | None, document_id: str | None) -> dict[str, Any]:
    validation = validate(status=status, language=language, document_id=document_id)
    catalog = load_catalog()
    create_reference_template(TEMPLATE_PATH)
    _normalise_docx(TEMPLATE_PATH)
    destination_root = output_root(catalog)
    destination_root.mkdir(parents=True, exist_ok=True)
    selected = selected_variants(catalog, status=status, language=language, document_id=document_id)
    outputs: list[dict[str, Any]] = []
    for document, current_language, variant in selected:
        aggregate_inputs: list[dict[str, str]] | None = None
        if variant.get("source") is None:
            source, aggregate_inputs = aggregate_manual_source(
                catalog,
                document,
                current_language,
            )
            source_path = source.path
            source_reference: str | None = None
            source_hash = hashlib.sha256(source.body.encode("utf-8")).hexdigest()
        else:
            source_path = safe_manual_path(str(variant["source"]))
            source = parse_manual_source(source_path)
            source_reference = source_path.relative_to(MANUAL_ROOT).as_posix()
            source_hash = sha256_file(source_path)
        output = destination_root / current_language / str(variant["output"])
        metrics = build_manual_docx(
            template=TEMPLATE_PATH,
            catalog=catalog,
            document_record=document,
            language=current_language,
            metadata=source.metadata,
            source_path=source_path,
            body=source.body,
            output=output,
        )
        _normalise_docx(output)
        audit = _audit_docx(output, expected_title=str(source.metadata["title"]))
        outputs.append(
            {
                "document_id": document["document_id"],
                "language": current_language,
                "status": variant["status"],
                "source": source_reference,
                "source_sha256": source_hash,
                "aggregate_sources": aggregate_inputs,
                "output": output.relative_to(destination_root).as_posix(),
                "output_sha256": sha256_file(output),
                "output_bytes": output.stat().st_size,
                "metrics": metrics,
                "audit": audit,
            }
        )
    shutil.copy2(CATALOG_PATH, destination_root / "source-catalog.json")
    build_manifest = {
        "schema_version": "pilot-assessment-document-build-v1",
        "product_version": catalog["product_version"],
        "document_set_version": catalog["document_set_version"],
        "release_channel": catalog["release_channel"],
        "release_label": catalog["release_label"],
        "user_acceptance": catalog["user_acceptance"],
        "build_status": status,
        "template": {
            "path": TEMPLATE_PATH.relative_to(MANUAL_ROOT).as_posix(),
            "sha256": sha256_file(TEMPLATE_PATH),
        },
        "toolchain": {
            "python_lock_sha256": sha256_file(TOOL_ROOT / "uv.lock"),
            "node_lock_sha256": sha256_file(TOOL_ROOT / "pnpm-lock.yaml"),
            "style_preset": "compact_reference_guide",
            "cover_pattern": "editorial_cover",
        },
        "outputs": outputs,
    }
    write_json(destination_root / "documentation-manifest.json", build_manifest)
    return {
        "validation": validation,
        "destination": str(destination_root),
        "template_sha256": build_manifest["template"]["sha256"],
        "outputs": outputs,
        "manifest_sha256": sha256_file(destination_root / "documentation-manifest.json"),
        "status": "PASS",
    }


def _arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    add_selection_arguments(parser)
    return parser.parse_args()


def main() -> int:
    args = _arguments()
    try:
        result = build(
            status=args.status,
            language=args.language,
            document_id=args.document_id,
        )
    except (DocumentationError, OSError, ValueError, zipfile.BadZipFile) as error:
        print(f"documentation build failed: {error}", file=sys.stderr)
        return 1
    print(result_payload(**result))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
